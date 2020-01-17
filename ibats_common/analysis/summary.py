#! /usr/bin/env python
# -*- coding:utf-8 -*-
"""
@author  : MG
@Time    : 19-5-9 上午9:53
@File    : summary.py
@contact : mmmaaaggg@163.com
@desc    : 
"""
import datetime
import itertools
import json
import logging
import os
from collections import defaultdict

import docx
import ffn
import numpy as np
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from ibats_utils.mess import open_file_with_system_app, date_2_str, datetime_2_str, split_chunk, iter_2_range, \
    format_2_str
from scipy.stats import anderson, normaltest

from ibats_common.analysis.plot import drawdown_plot, plot_rr_df, wave_hist, plot_scatter_matrix, plot_corr, \
    clean_cache, hist_n_rr, label_distribution, show_dl_accuracy
from ibats_common.analysis.plot_db import get_rr_with_md, show_trade, show_cash_and_margin
from ibats_common.backend.mess import get_report_folder_path
from ibats_common.backend.mess import get_stg_run_info
from ibats_common.common import RunMode, CalcMode

logger = logging.getLogger(__name__)
FFN_VERSION = ffn.__version__
STR_FORMAT_DATETIME_4_FILE_NAME = '%Y-%m-%d %H_%M_%S'
FORMAT_2_PERCENT = lambda x: f"{x * 100: .2f}%"
FORMAT_2_FLOAT2 = r"{0:.2f}"
FORMAT_2_FLOAT4 = r"{0:.4f}"
FORMAT_2_MONEY = r"{0:,.2f}"


def summary_md(md_df: pd.DataFrame, percentiles=[0.2, 1 / 3, 0.5, 2 / 3, 0.8],
               risk_free=0.03,
               close_key=None,
               enable_show_plot=True,
               enable_save_plot=False,
               name=None,
               func_kwargs_dic={},
               **kwargs,
               ):
    """
    汇总展示数据分析结果，同时以 dict 形式返回各项指标分析结果
    第一个返回值，df的各项分析结果
    第二个返回值，各个列的各项分析结果
    第三个返回值，相关文件输出路径或路径列表（当 enable_save_plot=True）
    :param md_df:
    :param risk_free:无风险收益率
    :param close_key:对哪些列的数据执行统计
    :param enable_show_plot: 展示plot
    :param enable_save_plot: 保存文件
    :param name:
    :param func_kwargs_dic:
    :param kwargs:
    :return:
    """
    columns = list(md_df.columns)
    logger.info('data columns: %s', columns)
    ret_dic, each_col_dic, file_path_dic = {}, defaultdict(dict), {}

    # logger.info('Description:')
    # df.describe(percentiles=percentiles)

    # 检查数据
    func_kwargs = func_kwargs_dic.setdefault("validation", {})
    trade_date_max_gap = func_kwargs.setdefault('trade_date_max_gap', 15)
    trade_date_s = pd.Series(md_df.index)
    days_2_next_trade_date_s = (trade_date_s.shift(-1) - trade_date_s).fillna(pd.Timedelta(days=0))
    days_2_next_trade_date_s.index = md_df.index
    over_gap_s = days_2_next_trade_date_s[days_2_next_trade_date_s.apply(lambda x: x.days > trade_date_max_gap)]
    warn_msg = [f"{date_2_str(trade_date)} 到下一个交易日间存在 {delta.days} 天没有行情数据，可能存在数据遗漏"
                for trade_date, delta in over_gap_s.items()]
    if len(warn_msg) > 0:
        ret_dic['warning'] = warn_msg

    func_kwargs = func_kwargs_dic.setdefault("rr_quantile", None)
    if func_kwargs is not None:
        col_name_list = func_kwargs.setdefault('columns', columns)
        quantile_df = md_df[col_name_list].to_returns().quantile(percentiles).rename(
            columns={_: _ + ' rr' for _ in col_name_list}).T
        ret_dic['rr_quantile'] = quantile_df

    # 获取统计数据
    stats = md_df.calc_stats()
    stats.set_riskfree_rate(risk_free)
    ret_dic['stats'] = stats
    enable_kwargs = {"enable_save_plot": enable_save_plot, "enable_show_plot": enable_show_plot,
                     "name": name}

    # scatter_matrix
    # diagonal，必须且只能在{'hist', 'kde'}中选择1个，
    # 'hist'表示直方图(Histogram plot),'kde'表示核密度估计(Kernel Density Estimation)
    # 该参数是scatter_matrix函数的关键参数
    file_path = plot_scatter_matrix(md_df, diagonal='kde', **enable_kwargs)
    if enable_save_plot:
        file_path_dic['scatter_matrix'] = file_path

    # stats.plot_correlation()
    file_path = plot_corr(md_df, **enable_kwargs)
    if enable_save_plot:
        file_path_dic['correlation'] = file_path

    # return rate plot 图
    func_kwargs = func_kwargs_dic.setdefault('rr', {})
    file_path = plot_rr_df(md_df, **func_kwargs, **enable_kwargs)
    if enable_save_plot:
        file_path_dic['rr'] = file_path

    # histgram 分布图
    func_kwargs = func_kwargs_dic.setdefault('hist', {})
    n_bins_dic, file_path = wave_hist(md_df, **func_kwargs, **enable_kwargs)
    ret_dic['hist'] = n_bins_dic
    if enable_save_plot:
        file_path_dic['hist'] = file_path

    # 回撤图
    func_kwargs = func_kwargs_dic.setdefault('drawdown', {})
    drawdown_df, file_path = drawdown_plot(md_df, perf_stats=stats, **func_kwargs, **enable_kwargs)
    ret_dic['drawdown'] = drawdown_df
    if enable_save_plot:
        file_path_dic['drawdown'] = file_path

    # 未来N日收益率分布
    func_kwargs = func_kwargs_dic.setdefault('hist_future_n_rr', {})
    tmp_dic, file_path = hist_n_rr(md_df, **func_kwargs, **enable_kwargs)
    ret_dic['hist_future_n_rr'] = tmp_dic
    if enable_save_plot:
        file_path_dic['hist_future_n_rr'] = file_path

    file_path_dic['label_distribution'] = defaultdict(dict)
    ret_dic['label_distribution'] = defaultdict(dict)
    noname_enable_kwargs = enable_kwargs.copy()
    del noname_enable_kwargs['name']
    for (n_day, col_name), quantile_df in tmp_dic['quantile_dic'].items():
        tmp_path_dic = file_path_dic['label_distribution'][(n_day, col_name)]
        distribution_dic = ret_dic['label_distribution'][(n_day, col_name)]
        col_count = quantile_df.shape[1]
        for n in range(col_count):
            max_rr = quantile_df.iloc[0, n]
            min_rr = quantile_df.iloc[1, col_count - n - 1]
            distribution_rate_df, file_path = label_distribution(
                md_df[close_key], min_rr=min_rr, max_rr=max_rr, max_future=n_day,
                name=f"{col_name}[{min_rr * 100:.2f}%~{max_rr * 100:.2f}%]", **noname_enable_kwargs)
            tmp_path_dic[(min_rr, max_rr)] = file_path
            distribution_dic[(min_rr, max_rr)] = distribution_rate_df

    # 单列分析
    stat_col_name_list = [close_key]
    stat_df = (md_df if stat_col_name_list is None else md_df[stat_col_name_list])
    for col_name, data in stat_df.items():
        data_df = md_df[[col_name]].dropna()
        data = data.dropna()
        data = data[~np.isinf(data)]
        logger.info("=" * 50 + ' %s ' + "=" * 50, col_name)
        result = normaltest(data)
        each_col_dic[col_name]['normaltest'] = result
        logger.info('%s %s', col_name, result)
        result = anderson(data)
        logger.info('%s %s', col_name, result)
        each_col_dic[col_name]['anderson'] = result
        stats = data.calc_perf_stats()
        # stats = ffn_stats[col_name]
        # 设置无风险收益率
        stats.set_riskfree_rate(risk_free)
        stats.display()
        each_col_dic[col_name]['stats'] = stats
        # plot
        n_bins_dic, file_path = wave_hist(
            md_df[[col_name]].dropna().to_returns().rename(columns={col_name: f'{col_name} rr'}),
            figure_4_each_col=True, name=col_name, **noname_enable_kwargs)
        each_col_dic[col_name]['hist'] = n_bins_dic
        if enable_save_plot:
            file_path_dic[f'{col_name} hist'] = file_path[0]
        rr = data_df.to_returns()
        each_col_dic[col_name]['rr_quantile'] = rr.quantile(
            [0.25, 0.33, 0.5, 0.66, 0.75]
        ).rename(columns={col_name: f"{col_name} rr"}).T

    return ret_dic, each_col_dic, file_path_dic


def summary_rr(df: pd.DataFrame, risk_free=0.03,
               figure_4_each_col=False,
               col_transfer_dic: (dict, None) = None,
               stat_col_name_list=None,
               col_name_list=None,
               enable_show_plot=True,
               enable_save_plot=False,
               name=None,
               stg_run_id=None,
               **kwargs
               ):
    """
    汇总展示数据分析结果，同时以 dict 形式返回各项指标分析结果
    第一个返回值，df的各项分析结果
    第二个返回值，各个列的各项分析结果
    第三个返回值，相关文件输出路径或路径列表（当 enable_save_plot=True）
    :param df:
    :param risk_free:无风险收益率
    :param figure_4_each_col:hist图使用，每一列显示单独一张图片
    :param col_transfer_dic:列转换方法
    :param stat_col_name_list:对哪些列的数据执行统计
    :param col_name_list:对哪些列的数据执行统计
    :param enable_show_plot: 展示plot
    :param enable_save_plot: 保存文件
    :param name:
    :param stg_run_id:
    :param kwargs:
    :return:
    """
    ret_dic, each_col_dic, file_path_dic = {}, defaultdict(dict), {}
    if df is None:
        logger.error("传入数据为空")
        return ret_dic, each_col_dic, file_path_dic
    columns = list(df.columns)
    logger.info('data columns: %s', columns)

    # 获取统计数据
    stats = df.calc_stats()
    stats.set_riskfree_rate(risk_free)
    ret_dic['stats'] = stats
    enable_kwargs = {"enable_save_plot": enable_save_plot, "enable_show_plot": enable_show_plot,
                     "name": name, "stg_run_id": stg_run_id}

    # scatter_matrix
    # diagonal，必须且只能在{'hist', 'kde'}中选择1个，
    # 'hist'表示直方图(Histogram plot),'kde'表示核密度估计(Kernel Density Estimation)
    # 该参数是scatter_matrix函数的关键参数
    file_path = plot_scatter_matrix(df, diagonal='kde', **enable_kwargs)
    if enable_save_plot:
        file_path_dic['scatter_matrix'] = file_path

    # stats.plot_correlation()
    file_path = plot_corr(df, **enable_kwargs)
    if enable_save_plot:
        file_path_dic['correlation'] = file_path

    # return rate plot 图
    file_path = plot_rr_df(df, **enable_kwargs)
    if enable_save_plot:
        file_path_dic['rr'] = file_path

    # histgram 分布图
    n_bins_dic, file_path = wave_hist(df, figure_4_each_col=figure_4_each_col, col_transfer_dic=col_transfer_dic,
                                      **enable_kwargs)
    ret_dic['hist'] = n_bins_dic
    if enable_save_plot:
        file_path_dic['hist'] = file_path

    # 回撤图
    drawdown_df, file_path = drawdown_plot(df, perf_stats=stats, col_name_list=col_name_list,
                                           **enable_kwargs)
    ret_dic['drawdown'] = drawdown_df
    if enable_save_plot:
        file_path_dic['drawdown'] = file_path

    # 单列分析
    stat_df = (df if stat_col_name_list is None else df[stat_col_name_list])
    for col_name, data in stat_df.items():
        data = data.dropna()
        data = data[~np.isinf(data)]
        logger.info("=" * 50 + ' %s ' + "=" * 50, col_name)
        result = normaltest(data)
        each_col_dic[col_name]['normaltest'] = result
        logger.info('%s %s', col_name, result)
        result = anderson(data)
        logger.info('%s %s', col_name, result)
        each_col_dic[col_name]['anderson'] = result
        stats = data.calc_perf_stats()
        # stats = ffn_stats[col_name]
        # 设置无风险收益率
        stats.set_riskfree_rate(risk_free)
        stats.display()

    return ret_dic, each_col_dic, file_path_dic


def df_2_table(doc, df, format_by_index=None, format_by_col=None, max_col_count=None,
               mark_top_n=None, mark_top_n_on_cols=None):
    """

    :param doc:
    :param df:
    :param format_by_index: 按索引格式化
    :param format_by_col: 按列格式化
    :param max_col_count: 每行最大列数（不包括索引）
    :param mark_top_n: 标记 top N
    :param mark_top_n_on_cols: 选择哪些列标记 top N，None 代表不标记
    :return:
    """
    if max_col_count is None:
        max_col_count = df.shape[1]

    if mark_top_n is not None:
        if mark_top_n_on_cols is not None:
            rank_df = df[mark_top_n_on_cols]
        else:
            rank_df = df
        rank_df = rank_df.rank(ascending=False)
        is_in_rank_df = rank_df <= mark_top_n
    else:
        is_in_rank_df = None

    for table_num, col_name_list in enumerate(split_chunk(list(df.columns), max_col_count)):
        if table_num > 0:
            # 如果是换行写入第二、三、四。。个表格，先打一个空行
            doc.add_paragraph('')

        sub_df = df[col_name_list]
        row_num, col_num = sub_df.shape
        t = doc.add_table(row_num + 1, col_num + 1)

        # write head
        # col_name_list = list(sub_df.columns)
        for j in range(col_num):
            # t.cell(0, j).text = df.columns[j]
            # paragraph = t.cell(0, j).add_paragraph()
            paragraph = t.cell(0, j + 1).paragraphs[0]
            paragraph.add_run(str(col_name_list[j])).bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # write head bg color
        for j in range(col_num + 1):
            # t.cell(0, j).text = df.columns[j]
            t.cell(0, j)._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="00A2E8"/>'.format(nsdecls('w'))))

        # format table style to be a grid
        t.style = 'TableGrid'

        # populate the table with the dataframe
        for i in range(row_num):
            index = sub_df.index[i]
            paragraph = t.cell(i + 1, 0).paragraphs[0]
            index_str = str(date_2_str(index))
            paragraph.add_run(index_str).bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if format_by_index is not None and index in format_by_index:
                format_row = format_by_index[index]
            else:
                format_row = None

            for j in range(col_num):
                col_name = col_name_list[j]
                if format_row is None and format_by_col is not None and col_name in format_by_col:
                    format_cell = format_by_col[col_name]
                else:
                    format_cell = format_row

                content = sub_df.values[i, j]
                if format_cell is None:
                    text = str(content)
                elif isinstance(format_cell, str):
                    text = str.format(format_cell, content)
                elif callable(format_cell):
                    text = format_cell(content)
                else:
                    raise ValueError('%s: %s 无效', index, format_cell)

                paragraph = t.cell(i + 1, j + 1).paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                try:
                    style = paragraph.add_run(text)
                    if is_in_rank_df is not None and col_name in is_in_rank_df and is_in_rank_df.loc[index, col_name]:
                        style.font.color.rgb = RGBColor(0xed, 0x1c, 0x24)
                        style.bold = True
                except TypeError as exp:
                    logger.exception('df.iloc[%d, %d] = df["%s", "%s"] = %s', i, j, index, col_name, text)
                    raise exp from exp

        for i in range(1, row_num + 1):
            for j in range(col_num + 1):
                if i % 2 == 0:
                    t.cell(i, j)._tc.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="A3D9EA"/>'.format(nsdecls('w'))))


def _test_df_2_table():
    df = pd.DataFrame({'a': list(range(10, 20)), 'b': list(range(10, 0, -1))})
    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    document.add_heading(f'测试使用', 1)
    df_2_table(document, df, mark_top_n_on_cols=df.columns, mark_top_n=3)
    # file_path = os.path.abspath(os.path.join(os.path.curdir, 'test.docx'))
    file_path = "/home/mg/github/code_mess/drl/drl_off_example/d3qn_replay_2019_08_25/output/2013-11-08/model/reports/test.docx"
    document.save(file_path)
    open_file_with_system_app(file_path)


def _test_summary_md():
    from ibats_common.example.data import load_data

    df = load_data('RB.csv').set_index('trade_date').drop('instrument_type', axis=1)
    df.index = pd.DatetimeIndex(df.index)
    col_transfer_dic = {
        'return': ['open', 'high', 'low', 'close', 'volume']
    }
    ret_dic, each_col_dic, file_path_dic = summary_md(
        df, enable_show_plot=True, enable_save_plot=False,
        func_kwargs_dic={
            "hist": {
                "figure_4_each_col": False,
                "col_transfer_dic": col_transfer_dic,
            },
            "drawdown": {
                "col_name_list": ['close'],
                "stat_col_name_list": ['close'],
            }
        })


def dic_2_table(doc, param_dic: dict, col_group_num=1,
                format_key=None, format_dic: (dict, None) = None, col_width=[1.75, 4.25]):
    # Highlight all cells limegreen (RGB 32CD32) if cell contains text "0.5"
    data_count = len(param_dic)
    if data_count == 0:
        return
    col_num_max = col_group_num * 2
    row_num_max = (data_count // col_group_num + 1) if data_count % col_group_num == 0 else \
        (data_count // col_group_num + 2)
    t = doc.add_table(row_num_max, col_num_max)

    # write head
    # col_name_list = list(sub_df.columns)
    for j in range(col_num_max):
        paragraph = t.cell(0, j).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if j % 2 == 0:
            paragraph.add_run('key').bold = True
        else:
            paragraph.add_run('value').bold = True

    # write head bg color
    for j in range(col_num_max):
        # t.cell(0, j).text = df.columns[j]
        t.cell(0, j)._tc.get_or_add_tcPr().append(
            parse_xml(r'<w:shd {} w:fill="00A2E8"/>'.format(nsdecls('w'))))

    # format table style to be a grid
    t.style = 'TableGrid'

    # write key value
    for num, (key, value) in enumerate(param_dic.items()):
        row_num, col_num = num // col_group_num + 1, num % col_group_num * 2

        paragraph = t.cell(row_num, col_num).paragraphs[0]
        # populate the table with the dataframe
        text = format_2_str(key, format_key)

        paragraph.add_run(text).bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if format_dic is not None and key in format_dic:
            format_cell = format_dic[key]
        else:
            format_cell = None

        text = format_2_str(value, format_cell)
        paragraph = t.cell(row_num, col_num + 1).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            style = paragraph.add_run(text)
        except TypeError as exp:
            logger.exception('param_dic[%s] = %s', key, text)
            raise exp from exp

    # set alternative color
    for i in range(1, row_num_max):
        for j in range(col_num_max):
            if i % 2 == 0:
                t.cell(i, j)._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="A3D9EA"/>'.format(nsdecls('w'))))

    # 如果出现多列，则将列的宽度 / 列数
    widths = (Inches(col_width[0] / col_group_num), Inches(col_width[1] / col_group_num))
    width_in_row = itertools.chain.from_iterable(itertools.repeat(widths, col_group_num))
    t.autofit = False
    # set col width
    # https://stackoverflow.com/questions/43051462/python-docx-how-to-set-cell-width-in-tables
    # It's not work
    # for row in t.rows:
    #     for cell, width in zip(row.cells, width_in_row):
    #         cell.width = width
    # it's work
    for idx, width in enumerate(width_in_row):
        t.columns[idx].width = width


def _test_dic_2_table():
    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = docx.shared.Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = docx.shared.RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    # 文件内容
    heading_count, sub_heading_count = 0, 0
    document.add_heading('测试使用', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')

    dic = {
        'from_date': str_2_date('2018-08-09'),
        'name': 'name test',
        'num': 12345,
        'float': 23456.789,
        'list': [1, 2, 3, 4],
        'long text': """bla bla bla（唧唧歪歪） 在英语中是“等等，之类的”的意思。
        当别人知道你要表达的意思时，用blablabla会比较方便，可要注意用的场合与人。在某种场合也形容某些人比较八卦。"""
    }
    file_path = 'test.docx'
    dic_2_table(document, dic, col_group_num=2, )
    document.save(file_path)
    open_file_with_system_app(file_path)
    return file_path


def summary_stg(stg_run_id=None):
    from ibats_common.analysis.plot_db import show_order, show_cash_and_margin, show_rr_with_md
    info = get_stg_run_info(stg_run_id)
    stg_run_id = info.stg_run_id

    data_dict, file_path = show_order(stg_run_id)
    df = show_cash_and_margin(stg_run_id)
    sum_df, symbol_rr_dic, save_file_path_dic = show_rr_with_md(stg_run_id)
    summary_rr(sum_df, figure_4_each_col=True, col_transfer_dic={'return': sum_df.columns})
    # for symbol, rr_df in symbol_rr_dic.items():
    #     col_transfer_dic = {'return': rr_df.columns}
    #     summary_rr(rr_df, figure_4_each_col=True, col_transfer_dic=col_transfer_dic)


def _test_summary_stg():
    stg_run_id = None
    summary_stg(stg_run_id)


def stats_df_2_docx_table(stats_df, document, format_axis='index', **kwargs):
    """
    将 stats_df 统计信息以表格形式写入 docx 文件中
    :param stats_df:
    :param document:
    :param format_axis: 'column' / 'index' 选择state是按 col 还是 index 进行格式化，默认是 col
    :param kwargs:
    :return:
    """

    format_by_index = {
        "total_return": FORMAT_2_PERCENT,
        "cagr": FORMAT_2_PERCENT,
        "mtd": FORMAT_2_PERCENT,
        "three_month": FORMAT_2_PERCENT,
        "six_month": FORMAT_2_PERCENT,
        "ytd": FORMAT_2_PERCENT,
        "one_year": FORMAT_2_PERCENT,
        "three_year": FORMAT_2_PERCENT,
        "five_year": FORMAT_2_PERCENT,
        "ten_year": FORMAT_2_PERCENT,
        "best_day": FORMAT_2_PERCENT,
        "worst_day": FORMAT_2_PERCENT,
        "best_month": FORMAT_2_PERCENT,
        "worst_month": FORMAT_2_PERCENT,
        "best_year": FORMAT_2_PERCENT,
        "worst_year": FORMAT_2_PERCENT,
        "avg_drawdown": FORMAT_2_PERCENT,
        "avg_up_month": FORMAT_2_PERCENT,
        "win_year_perc": FORMAT_2_PERCENT,
        "twelve_month_win_perc": FORMAT_2_PERCENT,
        "incep": FORMAT_2_PERCENT,
        "rf": FORMAT_2_PERCENT,
        "max_drawdown": FORMAT_2_PERCENT,
        "calmar": FORMAT_2_FLOAT2,
        "daily_sharpe": FORMAT_2_FLOAT2,
        "daily_sortino": FORMAT_2_FLOAT2,
        "daily_mean": FORMAT_2_FLOAT2,
        "daily_vol": FORMAT_2_FLOAT2,
        "daily_skew": FORMAT_2_FLOAT2,
        "daily_kurt": FORMAT_2_FLOAT2,
        "monthly_sharpe": FORMAT_2_FLOAT2,
        "monthly_sortino": FORMAT_2_FLOAT2,
        "monthly_mean": FORMAT_2_FLOAT2,
        "monthly_vol": FORMAT_2_FLOAT2,
        "monthly_skew": FORMAT_2_FLOAT2,
        "monthly_kurt": FORMAT_2_FLOAT2,
        "yearly_sharpe": FORMAT_2_FLOAT2,
        "yearly_sortino": FORMAT_2_FLOAT2,
        "yearly_mean": FORMAT_2_FLOAT2,
        "yearly_vol": FORMAT_2_FLOAT2,
        "yearly_skew": FORMAT_2_FLOAT2,
        "yearly_kurt": FORMAT_2_FLOAT2,
        "avg_drawdown_days": FORMAT_2_FLOAT2,
        "avg_down_month": FORMAT_2_FLOAT2,
        "start": date_2_str,
        "end": date_2_str,
    }
    df_2_table(document, stats_df,
               format_by_index=format_by_index if format_axis == 'index' else None,
               format_by_col=format_by_index if format_axis == 'column' else None,
               max_col_count=4, **kwargs)


def summary_stg_2_docx(stg_run_id=None, enable_save_plot=True, enable_show_plot=False, enable_clean_cache=True,
                       doc_file_path=None):
    """
    生成策略分析报告
    :param stg_run_id:
    :param enable_save_plot:
    :param enable_show_plot:
    :param enable_clean_cache:
    :param doc_file_path: 可以是目录 或 文件名路径
    :return:
    """
    info = get_stg_run_info(stg_run_id)
    stg_run_id = info.stg_run_id
    stg_name = info.stg_name
    run_mode = RunMode(info.run_mode)
    kwargs = {"enable_show_plot": enable_show_plot, "enable_save_plot": enable_save_plot,
              "run_mode": run_mode, "stg_run_id": stg_run_id}
    if run_mode == RunMode.Backtest_FixPercent:
        sum_df, symbol_rr_dic = get_rr_with_md(stg_run_id, compound_rr=True)
    else:
        sum_df, symbol_rr_dic = get_rr_with_md(stg_run_id, compound_rr=False)

    if sum_df is None or sum_df.shape[0] == 0:
        logger.warning('stg_run_id=%d 没有获取到 sum_df', stg_run_id)
        file_path = None
        return file_path

    ret_dic, each_col_dic, file_path_dic = summary_rr(sum_df, **kwargs)

    _, file_path = show_trade(**kwargs)
    file_path_dic['trade'] = file_path
    df, file_path = show_cash_and_margin(**kwargs)
    file_path_dic['cash_and_margin'] = file_path

    # 生成 docx 文档将所需变量
    heading_title = f'策略分析报告[{stg_run_id} {stg_name}] ' \
        f'{date_2_str(min(sum_df.index))} - {date_2_str(max(sum_df.index))} ({sum_df.shape[0]} days)'

    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = docx.shared.Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = docx.shared.RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    # 文件内容
    document.add_heading(heading_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')
    heading_count = 1
    document.add_heading(f'{heading_count}、策略回测收益曲线', 1)
    # 增加图片（此处使用相对位置）
    document.add_picture(file_path_dic['rr'])  # , width=docx.shared.Inches(1.25)
    heading_count += 1
    # 添加分页符
    document.add_page_break()

    document.add_heading(f'{heading_count}、策略回撤曲线', 1)
    document.add_picture(file_path_dic['drawdown'])
    heading_count += 1
    document.add_page_break()

    if run_mode != RunMode.Backtest_FixPercent:
        document.add_heading(f'{heading_count}、现金与仓位堆叠图', 1)
        document.add_picture(file_path_dic['cash_and_margin'])
        heading_count += 1
        document.add_page_break()

    document.add_heading(f'{heading_count}、散点图矩阵图（Scatter Matrix）', 1)
    document.add_picture(file_path_dic['scatter_matrix'])
    heading_count += 1
    document.add_page_break()

    document.add_heading(f'{heading_count}、相关性矩阵图（Correlation）', 1)
    document.add_picture(file_path_dic['correlation'])
    heading_count += 1
    document.add_page_break()

    document.add_heading(f'{heading_count}、绩效统计数据（Porformance stat）', 1)
    stats_df = ret_dic['stats'].stats
    stats_df_2_docx_table(stats_df, document)
    heading_count += 1
    document.add_page_break()

    # 交易记录
    document.add_heading(f'{heading_count}、买卖点记录', 1)
    document.add_picture(file_path_dic['trade'])
    heading_count += 1
    document.add_page_break()

    # 保存文件
    if doc_file_path is not None:
        if os.path.isdir(doc_file_path):
            folder_path, file_name = doc_file_path, ''
        else:
            folder_path, file_name = os.path.split(doc_file_path)
    else:
        folder_path, file_name = get_report_folder_path(stg_run_id), ''

    if folder_path != '' and not os.path.exists(folder_path):
        os.makedirs(folder_path)

    if file_name == '':
        try:
            calc_mode_str = CalcMode(json.loads(info.trade_agent_params_list)[0]['calc_mode']).name + " "
        except:
            calc_mode_str = " "

        # run_mode_str = run_mode.name + " "
        file_name = f"{stg_run_id} {stg_name} {calc_mode_str}" \
            f"{date_2_str(min(sum_df.index))} - {date_2_str(max(sum_df.index))} ({sum_df.shape[0]} days) " \
            f"{datetime_2_str(datetime.datetime.now(), STR_FORMAT_DATETIME_4_FILE_NAME)}.docx"
        file_path = os.path.join(folder_path, file_name)
    else:
        file_path = doc_file_path

    document.save(file_path)
    if enable_clean_cache:
        clean_cache()

    return file_path


def _test_summary_stg_2_docx(auto_open_file=True):
    stg_run_id = None
    file_path = summary_stg_2_docx(stg_run_id, enable_clean_cache=True)
    if auto_open_file and file_path is not None:
        open_file_with_system_app(file_path)


def summary_md_2_docx(md_df: pd.DataFrame, percentiles=[0.2, 0.33, 0.5, 0.66, 0.8],
                      risk_free=0.03,
                      close_key=None,
                      enable_show_plot=True,
                      enable_save_plot=False,
                      name=None,
                      func_kwargs_dic={},
                      enable_clean_cache=True,
                      ):
    """
    汇总展示数据分析结果，同时以 dict 形式返回各项指标分析结果
    第一个返回值，df的各项分析结果
    第二个返回值，各个列的各项分析结果
    :param md_df:
    :param percentiles:分为数信息
    :param risk_free:无风险收益率
    :param close_key:对哪些列的数据执行统计
    :param enable_show_plot:显示plot
    :param enable_save_plot:保存plot
    :param name:
    :param func_kwargs_dic:
    :param enable_clean_cache:
    :return:
    """
    ret_dic, each_col_dic, file_path_dic = summary_md(
        md_df, percentiles=percentiles, risk_free=risk_free, close_key=close_key,
        enable_show_plot=enable_show_plot, enable_save_plot=enable_save_plot, name=name,
        func_kwargs_dic=func_kwargs_dic)

    logger.debug('file_path_dic')
    for num, (k, v) in enumerate(file_path_dic.items(), start=1):
        if isinstance(v, dict):
            for num2, (k2, v2) in enumerate(v.items(), start=1):
                if isinstance(v2, dict):
                    for num3, (k3, v3) in enumerate(v2.items(), start=1):
                        logger.debug("%d.%d.%d) %s %s %s -> %s", num, num2, num3, k, k2, k3, v3)
                else:
                    logger.debug("%d.%d) %s %s -> %s", num, num2, k, k2, v2)
        else:
            logger.debug("%d) %s -> %s", num, k, v)

    # 生成 docx 文档将所需变量
    heading_title = f'数据分析报告 {date_2_str(min(md_df.index))} - {date_2_str(max(md_df.index))} ({md_df.shape[0]} days)'

    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = docx.shared.Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = docx.shared.RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    # 文件内容
    document.add_heading(heading_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')
    heading_count = 1
    if 'warning' in ret_dic:
        document.add_heading(f'{heading_count}、警告信息（Warning）', 1)
        warning_list = ret_dic['warning']
        p = document.add_paragraph('')
        for msg in warning_list:
            r = p.add_run(msg)
            # r.bold = True
            r.font.color.rgb = RGBColor(0xaf, 0x26, 0x26)
            p.add_run('\n')

        heading_count += 1
        document.add_page_break()

    if 'rr' in file_path_dic:
        document.add_heading(f'{heading_count}、行情曲线', 1)
        # 增加图片（此处使用相对位置）
        document.add_picture(file_path_dic['rr'])  # , width=docx.shared.Inches(1.25)
        heading_count += 1
        # 添加分页符
        document.add_page_break()

    if 'hist' in file_path_dic:
        document.add_heading(f'{heading_count}、Histgram 分布图', 1)
        document.add_picture(file_path_dic['hist'])
        heading_count += 1

    if 'rr_quantile' in ret_dic:
        document.add_heading(f'{heading_count}、分位数信息（Quantile）', 1)
        rr_quantile_df = ret_dic['rr_quantile']
        format_by_col = {_: FORMAT_2_PERCENT for _ in rr_quantile_df.columns}
        df_2_table(document, rr_quantile_df, format_by_col=format_by_col, max_col_count=5)
        heading_count += 1
        document.add_page_break()

    if 'hist_future_n_rr' in file_path_dic:
        document.add_heading(f'{heading_count}、未来N日收益率最高最低值分布图', 1)
        quantile_dic = ret_dic['hist_future_n_rr']['quantile_dic']
        for num, ((n_day, col_name), file_path) in enumerate(file_path_dic['hist_future_n_rr'].items(), start=1):
            document.add_heading(f'{heading_count}.{num}) 未来 {n_day} 日 {col_name} 收益率最高最低值分布图', 2)
            document.add_picture(file_path)
            document.add_heading(f'{heading_count}.{num}.1) 分位数信息', 3)
            data_df = quantile_dic[(n_day, col_name)]
            df_2_table(document, data_df, format_by_index={_: FORMAT_2_PERCENT for _ in data_df.index})
            document.add_heading(f'{heading_count}.{num}.2) 三分类标签分布比例', 3)
            for (min_pct, max_pct), distribution_rate_df in ret_dic['label_distribution'][(n_day, col_name)].items():
                file_path = file_path_dic['label_distribution'][(n_day, col_name)][(min_pct, max_pct)]
                document.add_picture(file_path)
                distribution_rate_df.rename(
                    columns={1: f'1 under {min_pct * 100:.2f}%', 2: f'2 over {max_pct * 100:.2f}%'}, inplace=True)
                df_2_table(document, distribution_rate_df,
                           format_by_index={_: FORMAT_2_PERCENT for _ in distribution_rate_df.index})

            document.add_page_break()

        heading_count += 1

    if 'drawdown' in file_path_dic:
        document.add_heading(f'{heading_count}、行情回撤曲线', 1)
        document.add_picture(file_path_dic['drawdown'])
        heading_count += 1
        document.add_page_break()

    if 'scatter_matrix' in file_path_dic:
        document.add_heading(f'{heading_count}、散点图矩阵图（Scatter Matrix）', 1)
        document.add_picture(file_path_dic['scatter_matrix'])
        heading_count += 1
        document.add_page_break()

    if 'correlation' in file_path_dic:
        document.add_heading(f'{heading_count}、相关性矩阵图（Correlation）', 1)
        document.add_picture(file_path_dic['correlation'])
        heading_count += 1
        document.add_page_break()

    if 'stats' in file_path_dic:
        document.add_heading(f'{heading_count}、绩效统计数据（Porformance stat）', 1)
        stats_df = ret_dic['stats'].stats
        stats_df_2_docx_table(stats_df, document)
        heading_count += 1
        document.add_page_break()

    # 逐列分析
    col_name = close_key
    # 文件内容
    heading_title = f'{col_name} 数据分析结果'
    document.add_heading(heading_title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')
    heading_count = 1

    if 'rr_quantile' in ret_dic:
        document.add_heading(f'{heading_count}、分位数信息（Quantile）', 1)
        data_df = each_col_dic[col_name]['rr_quantile']
        format_by_col = {_: FORMAT_2_PERCENT for _ in data_df.columns}
        df_2_table(document, data_df, format_by_col=format_by_col, max_col_count=5)
        heading_count += 1

    if f'{col_name} hist' in file_path_dic:
        document.add_heading(f'{heading_count}、Histgram 分布图', 1)
        document.add_picture(file_path_dic[f'{col_name} hist'])
        heading_count += 1
        document.add_page_break()

    # 保存文件
    file_name = f"MD{' ' if name is None else ' ' + name} " \
        f"{date_2_str(min(md_df.index))} - {date_2_str(max(md_df.index))} " \
        f"({md_df.shape[0]} days) {datetime_2_str(datetime.datetime.now(), STR_FORMAT_DATETIME_4_FILE_NAME)}.docx"
    file_path = os.path.join(get_report_folder_path(), file_name)
    document.save(file_path)
    if enable_clean_cache:
        clean_cache()

    return file_path


def _test_summary_md_2_docx(auto_open_file=True):
    from ibats_common.example.data import load_data
    instrument_type = 'RU'  # 'RB' 'RU'
    file_name = f"{instrument_type}.csv"

    factor_df = load_data(file_name).set_index('trade_date').drop('instrument_type', axis=1)
    factor_df.index = pd.DatetimeIndex(factor_df.index)
    column_list_oraginal = list(factor_df.columns)
    ohlcav_col_name_list = ["open", "high", "low", "close", "amount", "volume"]

    from ibats_common.backend.factor import get_factor
    from ibats_common.example.data import get_trade_date_series
    from ibats_common.example.data import get_delivery_date_series
    factor_df = get_factor(factor_df, ohlcav_col_name_list=ohlcav_col_name_list,
                           trade_date_series=get_trade_date_series(),
                           delivery_date_series=get_delivery_date_series(instrument_type))

    col_transfer_dic = {
        'return': ['open', 'high', 'low', 'close', 'volume']
    }
    file_path = summary_md_2_docx(
        factor_df, enable_show_plot=False, enable_save_plot=True, close_key='close', name=instrument_type,
        func_kwargs_dic={
            "hist": {
                "figure_4_each_col": False,
                "columns": column_list_oraginal,
                "col_transfer_dic": col_transfer_dic,
            },
            "drawdown": {
                "col_name_list": ['close'],
            },
            "rr": {
                "col_name_list": ['close'],
            },
            "hist_future_n_rr": {
                'n_days': [3, 5],
                "columns": ['close'],
            },
            "rr_quantile": {'columns': ['close']},
            "validation": {'trade_date_max_gap': 10},
            # "": {},
        })
    if auto_open_file:
        open_file_with_system_app(file_path)


def summary_release_2_docx(title, img_meta_dic_list, stg_run_id=None, enable_clean_cache=True):
    """
    生成 预测成功率趋势报告
    :param title:
    :param img_meta_dic_list:
    :param stg_run_id:
    :param enable_clean_cache:
    :return:
    """
    logger.debug('生成报告开始')
    # 生成 docx 文件
    document = docx.Document()
    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')
    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = docx.shared.Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = docx.shared.RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '微软雅黑')

    # 文件内容
    document.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('')
    document.add_paragraph('')
    heading_size = 1
    for num, info_dic in enumerate(img_meta_dic_list, start=1):
        trade_date_last_train = info_dic['trade_date_last_train']
        trade_date_end = info_dic['trade_date_end']
        document.add_heading(
            f"{num}、{date_2_str(trade_date_last_train)} - {date_2_str(trade_date_end)}", heading_size)
        split_point_list = info_dic['split_point_list']
        if split_point_list is None:
            p = document.add_paragraph(f"{num}.1） 日期区间段1个:\n")
            p.add_run(f'\t1) {date_2_str(trade_date_last_train)} ~ {date_2_str(trade_date_end)}\n')
        else:
            p = document.add_paragraph(f"{num}.1） 日期区间段{len(split_point_list) - 1}个:\n")
            for num2, (point1, point2) in enumerate(
                    iter_2_range(split_point_list, has_left_outer=False, has_right_outer=False), start=1):
                p.add_run(f'\t{num2}) {date_2_str(point1)} ~ {date_2_str(point2)}\n')

        document.add_paragraph(f"{num}.2） 模型路径：\n\t{info_dic['module_file_path']}")
        document.add_paragraph(f"{num}.3） 取样状态(random_state)：\n\t{info_dic['predict_test_random_state']}")
        document.add_paragraph(f"{num}.4） 展示数据长度：\n\t{info_dic['in_range_count']}")
        document.add_paragraph(f"{num}.5） 预测准确率趋势图：")
        document.add_picture(info_dic['img_file_path'])

    file_name = f"{title}.docx"
    file_path = os.path.join(get_report_folder_path(stg_run_id), file_name)
    document.save(file_path)
    if enable_clean_cache:
        clean_cache()
    logger.debug('生成报告结束。%s', file_path)
    return file_path


def _test_summary_release_2_docx():
    real_ys, pred_ys = np.random.randint(1, 3, size=100), np.random.randint(1, 3, size=100)
    date_arr = pd.date_range(pd.to_datetime('2018-01-01'),
                             pd.to_datetime('2018-01-01') + pd.Timedelta(days=99))
    date_index = pd.DatetimeIndex(date_arr)
    close_df = pd.DataFrame({'close': np.sin(np.linspace(0, 10, 100))}, index=date_index)

    split_point_list = np.concatenate((np.random.randint(1, len(date_arr) - 1, size=10), [1, len(date_arr) - 1]))
    split_point_list.sort()
    split_point_list = date_arr[split_point_list]
    base_line_list = [0.3, 0.6]
    img_meta_dic_list = []
    img_file_path = show_dl_accuracy(real_ys, pred_ys, close_df, split_point_list, base_line_list, show_moving_avg=True)
    img_meta_dic_list.append({
        'img_file_path': img_file_path,
        'trade_date_last_train': pd.to_datetime('2018-01-01'),
        'module_file_path':
            "/home/mg/github/IBATS_Common/ibats_common/tf_saves_2019-06-25_08_13_36/model_tfls/2012-12-31",
        'predict_test_random_state': 1,
        'split_point_list': split_point_list,
        'in_range_count': close_df.shape[0],
        'trade_date_end': pd.to_datetime('2018-01-01') + pd.Timedelta(days=99),
    })
    img_file_path = show_dl_accuracy(real_ys, pred_ys, close_df, split_point_list, base_line_list,
                                     show_moving_avg=False)
    img_meta_dic_list.append({
        'img_file_path': img_file_path,
        'trade_date_last_train': pd.to_datetime('2018-01-01'),
        'module_file_path':
            "/home/mg/github/IBATS_Common/ibats_common/tf_saves_2019-06-25_08_13_36/model_tfls/2012-12-31",
        'predict_test_random_state': 1,
        'split_point_list': split_point_list,
        'in_range_count': close_df.shape[0],
        'trade_date_end': pd.to_datetime('2018-01-01') + pd.Timedelta(days=99),
    })
    file_path = summary_release_2_docx('test_only', img_meta_dic_list)
    open_file_with_system_app(file_path)


if __name__ == "__main__":
    pass
    # _test_summary_md()
    # _test_summary_stg()
    # _test_summary_stg_2_docx()
    # _test_summary_md_2_docx()
    # _test_summary_release_2_docx()
    _test_df_2_table()
